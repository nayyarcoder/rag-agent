import pytest
import tempfile
import os
from pathlib import Path
from ingestion import DocumentIngester

@pytest.fixture
def chunk_size():
    """Fixture to define the chunk size for testing."""
    return 1000

@pytest.fixture
def ingester(chunk_size):
    """Fixture to create a DocumentIngester instance with specified chunk size."""
    return DocumentIngester(chunk_size=chunk_size)

@pytest.fixture
def test_dir():
    """Fixture to create and clean up a temporary directory."""
    temp_dir = tempfile.mkdtemp()
    yield temp_dir
    import shutil
    shutil.rmtree(temp_dir)

def create_test_docx(test_dir, content="Test content for Word document"):
    """Create a test .docx file."""
    import docx
    doc_path = Path(test_dir) / "test.docx"
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(str(doc_path))
    return str(doc_path)

def create_test_txt(test_dir, content="Test content for text file"):
    """Create a test .txt file."""
    txt_path = Path(test_dir) / "test.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)
    return str(txt_path)

def test_docx_loading(ingester, test_dir):
    """Test loading a .docx file."""
    test_content = "This is a test Word document.\nIt contains multiple lines.\nTesting document loading."
    docx_path = create_test_docx(test_dir, test_content)
    
    # Test document loading
    documents = ingester.load_document(docx_path)
    
    # Verify documents were loaded
    assert len(documents) > 0, "No documents were loaded"
    
    # Verify content is present
    combined_content = " ".join([doc.page_content for doc in documents])
    assert "test Word document" in combined_content, "Expected content not found in loaded document"
    
    # Verify metadata
    assert all(doc.metadata.get("source") == docx_path for doc in documents), \
        "Source metadata not properly set"

def test_txt_loading(ingester, test_dir):
    """Test loading a .txt file for comparison."""
    test_content = "This is a test text file.\nIt contains multiple lines.\nTesting document loading."
    txt_path = create_test_txt(test_dir, test_content)
    
    # Test document loading
    documents = ingester.load_document(txt_path)
    
    # Verify documents were loaded
    assert len(documents) > 0, "No documents were loaded"
    
    # Verify content is present
    assert documents[0].page_content.strip() == test_content.strip(), \
        "Text content does not match input"
    
    # Verify metadata
    assert documents[0].metadata.get("source") == txt_path, \
        "Source metadata not properly set"

def test_document_processing(ingester, test_dir, chunk_size):
    """Test document chunking with Word documents."""
    # Create a longer document to test chunking
    long_content = "This is a test Word document.\n" * 20
    docx_path = create_test_docx(test_dir, long_content)
    
    # Load and process documents
    documents = ingester.load_document(docx_path)
    chunks = ingester.process_documents(documents)
    
    # Verify chunking
    assert len(chunks) > 0, "No chunks were created"
    assert all(len(chunk.page_content) <= chunk_size 
              for chunk in chunks), "Chunks exceed maximum size"
    
    # Additional verification of chunk contents
    total_content = "".join(chunk.page_content for chunk in chunks)
    assert "test Word document" in total_content, "Content was lost during chunking"
    
    # Verify chunk overlap
    if len(chunks) > 1:
        # Get the end of first chunk and start of second chunk
        first_chunk_end = chunks[0].page_content[-50:]  # Last 50 chars of first chunk
        second_chunk_start = chunks[1].page_content[:50]  # First 50 chars of second chunk
        # There should be some overlap between chunks
        assert any(word in second_chunk_start for word in first_chunk_end.split()), \
            "No overlap found between chunks"

def test_unsupported_file(ingester, test_dir):
    """Test handling of unsupported file types."""
    unsupported_path = Path(test_dir) / "test.xyz"
    with open(unsupported_path, "w") as f:
        f.write("test content")
        
    documents = ingester.load_document(str(unsupported_path))
    assert len(documents) == 0, "Unsupported file type should return empty list"

def test_docx_with_formatting(ingester, test_dir):
    """Test loading a .docx file with various formatting."""
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    # Create a document with rich formatting
    doc_path = Path(test_dir) / "formatted.docx"
    doc = docx.Document()
    
    # Add title
    title = doc.add_heading('Test Document', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add normal paragraph
    p1 = doc.add_paragraph('This is a normal paragraph with ')
    p1.add_run('bold').bold = True
    p1.add_run(' and ')
    p1.add_run('italic').italic = True
    p1.add_run(' text.')
    
    # Add bullet points
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    
    # Save the document
    doc.save(str(doc_path))
    
    # Test document loading
    documents = ingester.load_document(str(doc_path))
    
    # Verify content was loaded and formatting was preserved as text
    combined_content = " ".join([doc.page_content for doc in documents])
    assert "Test Document" in combined_content, "Title not found in content"
    assert "bullet point" in combined_content, "Bullet points not found in content" 